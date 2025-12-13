from typing import List
from analyzer import ChapterSummary, ImportantDetail
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from hebrew_numbers import int_to_gematria

class ImportantDetailSummary:
    name: str
    description: List[str]
    references: List[str]
    
    def __init__(self, name: str, description: List[str], references: List[str]):
        self.name = name
        self.description = description
        self.references = references
    
    def set(self, reference: str, description: str):
        if description not in self.description:
            self.description.append(description)
        if reference not in self.references:
            self.references.append(reference)

class Summarizer:
    doc: None
    characters: dict[str, ImportantDetailSummary] = {}
    places: dict[str, ImportantDetailSummary] = {}
    
    def __init__(self):
        self.characters = {}
        self.places = {}
        self.doc = Document()
        section = self.doc.sections[0]
        section.right_to_left = True
        
        title = self.doc.add_heading('מלכים', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_heading('עלילה', 1)  
    
    def set_summary(self, chapter: int, summary: ChapterSummary):            
        text = f'({int_to_gematria(chapter)}) {summary.synopsis}'            
        self.doc.add_paragraph(text)        
        for character in summary.characters:
            self.add_character(summary.chapter, character)
        for place in summary.places:
            self.add_place(summary.chapter,place)        
    
    def add_character(self, reference:str, character: ImportantDetail):
        if character.name not in self.characters:
            self.characters[character.name] = ImportantDetailSummary(
                name=character.name,
                description=[character.description],
                references=[reference]
            )
        else:
            existing = self.characters[character.name]
            existing.set(reference, character.description)
        
    def add_place(self, reference:str, place: ImportantDetail):
        if place.name not in self.places:
            self.places[place.name] = ImportantDetailSummary(
                name=place.name,
                description=[place.description],
                references=[reference]
            )
        else:
            existing = self.places[place.name]
            existing.set(reference, place.description)
    
    def save(self, filename='kings_summary.docx'):
        self.characters = dict(sorted(self.characters.items(), key=lambda x: x[0]))
        self.places = dict(sorted(self.places.items(), key=lambda x: x[0]))
        
        self.doc.add_heading('דמויות', 1)
        for character in self.characters.values():
            self.add_important_detail_to_doc(character)
        
        self.doc.add_heading('מקומות', 1)
        for place in self.places.values():
            self.add_important_detail_to_doc(place)
        
        self.doc.save(filename)
        print(f"Summary saved to {filename}")
        
    def add_important_detail_to_doc(self, detail: ImportantDetailSummary):
        # Add name in bold with colon
        p = self.doc.add_paragraph()
        run = p.add_run(f'{detail.name}: ')
        run.bold = True

        # Add descriptions
        for desc in detail.description:
            p.add_run(desc)
            if not desc.endswith('.'):
                p.add_run('. ')
            else:
                p.add_run(' ')

        p.add_run(' ')                        
        sorted_refs = sorted(detail.references)
        refs_text = ', '.join(sorted_refs)
        p.add_run(f"({refs_text})")        
                